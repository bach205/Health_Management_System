import shutil
import os
import asyncio
from pathlib import Path
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangchainDocument
from src.server.model.documentsModel import save_documents_to_database,save_chunks_into_chroma, delete_document_by_id, get_document_location_by_id,do_query_for_all_documents
from datetime import datetime

UPLOAD_FOLDER = Path(__file__).parent.parent.parent.parent / "documents"

async def save_documents(file,user_id):
    try:
        time_str = datetime.now().strftime("%Y%m%d%H%M%S")
        file.filename = time_str + "_" + file.filename
        # Tạo đường dẫn tuyệt đối để lưu file
        file_location = os.path.join(UPLOAD_FOLDER, file.filename)
        
        # Lưu file, wb la method write + binary (lam viec voi file nhi phan)
        with open(file_location, "wb") as buffer:
            #copy file object of file sang file object of buffer, (chung quy lai la ghi noi dung file vao buffer)
            shutil.copyfileobj(file.file, buffer)
        
        #save documents to database
    
        inserted_id = await save_documents_to_database(file,file_location,user_id)


        asyncio.create_task(asyncio.to_thread(index_documents, file_location, inserted_id))
    except Exception as e:
        print(e)
        raise Exception(e)
    return file

def load_docx_to_text(path):

    doc = Document(path)
    text = "\n".join([para.text for para in doc.paragraphs])
    return [text]

def load_docx_to_documents(path: str, doc_id=None):
    doc = Document(path)
    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip() != ""])
    metadata = {
        "source": os.path.basename(path),
    }
    if doc_id is not None:
        metadata["doc_id"] = doc_id
    return [LangchainDocument(page_content=text, metadata=metadata)]

def chunk_texts(texts):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=300,
        chunk_overlap=100
    )
    chunks = []
    for text in texts:
        chunks.extend(splitter.split_text(text))
    return chunks

def chunk_docs(docs):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=300,
        chunk_overlap=100
    )
    chunks = []
    chunks.extend(splitter.split_documents(docs))
    return chunks

#load documents and chunk them into chunks and save them into chroma
def index_documents(path:str, doc_id=None):
    docs = load_docx_to_documents(path, doc_id)
    chunks = chunk_docs(docs)
    save_chunks_into_chroma(chunks)
    print("Done")

async def delete_documents(id: int):
    try:
        await delete_document_by_id(id)
        # Xóa file vật lý trên ổ đĩa
        file_location, _ = await get_document_location_by_id(id)
        if file_location and os.path.exists(file_location):
            os.remove(file_location)
        return True
    except Exception as e:
        print(e)
        raise Exception(e)

async def get_document_file(id: int):
    file_location, file_name = await get_document_location_by_id(id)
    if file_location and os.path.exists(file_location):
        return file_location, file_name
    return None, None

async def query_for_all_documents():
    return await do_query_for_all_documents()